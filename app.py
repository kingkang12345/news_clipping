import streamlit as st
import re


# ✅ 무조건 첫 Streamlit 명령어
st.set_page_config(
    page_title="PwC 뉴스 분석기",
    page_icon="📊",
    layout="wide",
)



from datetime import datetime, timedelta, timezone
import os
from PIL import Image
import docx
from docx.shared import Pt, RGBColor, Inches
import io
from urllib.parse import urlparse
from googlenews import GoogleNews
from news_ai import (
    collect_news,
    filter_valid_press,
    filter_excluded_news,
    group_and_select_news,
    evaluate_importance,
)

# Import centralized configuration
from config import (
    COMPANY_CATEGORIES,
    COMPANY_KEYWORD_MAP,
    TRUSTED_PRESS_ALIASES,
    ADDITIONAL_PRESS_ALIASES,
    SYSTEM_PROMPT_1,
    SYSTEM_PROMPT_2,
    SYSTEM_PROMPT_3,
    EXCLUSION_CRITERIA,
    DUPLICATE_HANDLING,
    SELECTION_CRITERIA, 
    GPT_MODELS,
    DEFAULT_GPT_MODEL,
    # 새로 추가되는 회사별 기준들
    COMPANY_ADDITIONAL_EXCLUSION_CRITERIA,
    COMPANY_ADDITIONAL_DUPLICATE_HANDLING,
    COMPANY_ADDITIONAL_SELECTION_CRITERIA
)

# 한국 시간대(KST) 정의
KST = timezone(timedelta(hours=9))


def format_date(date_str):
    """Format date to MM/DD format with proper timezone handling"""
    try:
        # Try YYYY-MM-DD format
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        return date_obj.strftime('%m/%d')
    except Exception:
        try:
            # Try GMT format and convert to KST
            date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
            # Convert UTC to KST (add 9 hours)
            date_obj_kst = date_obj + timedelta(hours=9)
            return date_obj_kst.strftime('%m/%d')
        except Exception:
            try:
                # Try GMT format without timezone indicator
                date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S GMT')
                # Convert UTC to KST (add 9 hours)
                date_obj_kst = date_obj + timedelta(hours=9)
                return date_obj_kst.strftime('%m/%d')
            except Exception:
                # Return original if parsing fails
                return date_str if date_str else '날짜 정보 없음'

# 회사별 추가 기준을 적용하는 함수들
def get_enhanced_exclusion_criteria(companies):
    """회사별 제외 기준을 추가한 프롬프트 반환 (여러 회사 지원)"""
    base_criteria = EXCLUSION_CRITERIA
    
    # companies가 문자열이면 리스트로 변환
    if isinstance(companies, str):
        companies = [companies]
    
    # 선택된 모든 회사의 추가 기준을 합침
    all_additional_criteria = ""
    for company in companies:
        additional_criteria = COMPANY_ADDITIONAL_EXCLUSION_CRITERIA.get(company, "")
        if additional_criteria:
            all_additional_criteria += additional_criteria
    
    return base_criteria + all_additional_criteria

def get_enhanced_duplicate_handling(companies):
    """회사별 중복 처리 기준을 추가한 프롬프트 반환 (여러 회사 지원)"""
    base_criteria = DUPLICATE_HANDLING
    
    # companies가 문자열이면 리스트로 변환
    if isinstance(companies, str):
        companies = [companies]
    
    # 선택된 모든 회사의 추가 기준을 합침
    all_additional_criteria = ""
    for company in companies:
        additional_criteria = COMPANY_ADDITIONAL_DUPLICATE_HANDLING.get(company, "")
        if additional_criteria:
            all_additional_criteria += additional_criteria
    
    return base_criteria + all_additional_criteria

def get_enhanced_selection_criteria(companies):
    """회사별 선택 기준을 추가한 프롬프트 반환 (여러 회사 지원)"""
    base_criteria = SELECTION_CRITERIA
    
    # companies가 문자열이면 리스트로 변환
    if isinstance(companies, str):
        companies = [companies]
    
    # 선택된 모든 회사의 추가 기준을 합침
    all_additional_criteria = ""
    for company in companies:
        additional_criteria = COMPANY_ADDITIONAL_SELECTION_CRITERIA.get(company, "")
        if additional_criteria:
            all_additional_criteria += additional_criteria
    
    return base_criteria + all_additional_criteria
            
# 워드 파일 생성 함수
def create_word_document(keyword, final_selection, analysis=""):
    # 새 워드 문서 생성
    doc = docx.Document()
    
    # 제목 스타일 설정
    title = doc.add_heading(f'PwC 뉴스 분석 보고서: {keyword}', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(208, 74, 2)  # PwC 오렌지 색상
    
    # 분석 요약 추가
    if analysis:
        doc.add_heading('회계법인 관점의 분석 결과', level=1)
        doc.add_paragraph(analysis)
    
    # 선별된 주요 뉴스 추가
    doc.add_heading('선별된 주요 뉴스', level=1)
    
    for i, news in enumerate(final_selection):
        p = doc.add_paragraph()
        p.add_run(f"{i+1}. {news['title']}").bold = True
        
        # 날짜 정보 추가
        date_str = news.get('date', '날짜 정보 없음')
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f"날짜: {date_str}").italic = True
        
        # 선정 사유 추가
        reason = news.get('reason', '')
        if reason:
            doc.add_paragraph(f"선정 사유: {reason}")
        
        # 키워드 추가
        keywords = news.get('keywords', [])
        if keywords:
            doc.add_paragraph(f"키워드: {', '.join(keywords)}")
        
        # 관련 계열사 추가
        affiliates = news.get('affiliates', [])
        if affiliates:
            doc.add_paragraph(f"관련 계열사: {', '.join(affiliates)}")
        
        # 언론사 추가
        press = news.get('press', '알 수 없음')
        doc.add_paragraph(f"언론사: {press}")
        
        # URL 추가
        url = news.get('url', '')
        if url:
            doc.add_paragraph(f"출처: {url}")
        
        # 구분선 추가
        if i < len(final_selection) - 1:
            doc.add_paragraph("").add_run().add_break()
    
    # 날짜 및 푸터 추가
    current_date = datetime.now().strftime("%Y년 %m월 %d일")
    doc.add_paragraph(f"\n보고서 생성일: {current_date}")
    doc.add_paragraph("© 2024 PwC 뉴스 분석기 | 회계법인 관점의 뉴스 분석 도구")
    
    return doc

# BytesIO 객체로 워드 문서 저장
def get_binary_file_downloader_html(doc, file_name):
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# 커스텀 CSS
st.markdown("""
<style>
    .title-container {
        display: flex;
        align-items: center;
        gap: 20px;
        margin-bottom: 20px;
    }
    .main-title {
        color: #d04a02;
        font-size: 2.5rem;
        font-weight: 700;
    }
    .news-card {
        background-color: #f9f9f9;
        border-radius: 10px;
        padding: 15px;
        margin-bottom: 15px;
        border-left: 4px solid #d04a02;
    }
    .news-title {
        font-weight: 600;
        font-size: 1.1rem;
    }
    .news-url {
        color: #666;
        font-size: 0.9rem;
    }
    .news-date {
        color: #666;
        font-size: 0.9rem;
        font-style: italic;
        margin-top: 5px;
    }
    .analysis-box {
        background-color: #f5f5ff;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border-left: 4px solid #d04a02;
    }
    .subtitle {
        color: #dc582a;
        font-size: 1.3rem;
        font-weight: 600;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    .download-box {
        background-color: #eaf7f0;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border-left: 4px solid #00a36e;
        text-align: center;
    }
    .analysis-section {
        background-color: #f8f9fa;
        border-left: 4px solid #d04a02;
        padding: 20px;
        margin: 10px 0;
        border-radius: 5px;
    }
    .selected-news {
        border-left: 4px solid #0077b6;
        padding: 15px;
        margin: 10px 0;
        background-color: #f0f8ff;
        border-radius: 5px;
    }
    .excluded-news {
        color: #666;
        padding: 5px 0;
        margin: 5px 0;
        font-size: 0.9em;
    }
    .news-meta {
        color: #666;
        font-size: 0.9em;
        margin: 3px 0;
    }
    .selection-reason {
        color: #666;
        margin: 5px 0;
        font-size: 0.95em;
    }
    .keywords {
        color: #666;
        font-size: 0.9em;
        margin: 5px 0;
    }
    .affiliates {
        color: #666;
        font-size: 0.9em;
        margin: 5px 0;
    }
    .news-url {
        color: #0077b6;
        font-size: 0.9em;
        margin: 5px 0;
        word-break: break-all;
    }
    .news-title-large {
        font-size: 1.2em;
        font-weight: 600;
        color: #000;
        margin-bottom: 8px;
        line-height: 1.4;
    }
    .news-url {
        color: #0077b6;
        font-size: 0.9em;
        margin: 5px 0 10px 0;
        word-break: break-all;
    }
    .news-summary {
        color: #444;
        font-size: 0.95em;
        margin: 10px 0;
        line-height: 1.4;
    }
    .selection-reason {
        color: #666;
        font-size: 0.95em;
        margin: 10px 0;
        line-height: 1.4;
    }
    .importance-high {
        color: #d04a02;
        font-weight: 700;
        margin: 5px 0;
    }
    .importance-medium {
        color: #0077b6;
        font-weight: 700;
        margin: 5px 0;
    }
    .group-indices {
        color: #666;
        font-size: 0.9em;
    }
    .group-selected {
        color: #00a36e;
        font-weight: 600;
    }
    .group-reason {
        color: #666;
        font-size: 0.9em;
        margin-top: 5px;
    }
    .not-selected-news {
        color: #666;
        padding: 5px 0;
        margin: 5px 0;
        font-size: 0.9em;
    }
    .importance-low {
        color: #666;
        font-weight: 700;
        margin: 5px 0;
    }
    .not-selected-reason {
        color: #666;
        margin: 5px 0;
        font-size: 0.95em;
    }
    .email-preview {
        background-color: white;
        border: 1px solid #ddd;
        border-radius: 5px;
        padding: 20px;
        margin: 20px 0;
        overflow-y: auto;
        max-height: 500px;
    }
    .copy-button {
        background-color: #d04a02;
        color: white;
        padding: 10px 20px;
        border: none;
        border-radius: 5px;
        cursor: pointer;
        margin: 10px 0;
    }
    .copy-button:hover {
        background-color: #b33d00;
    }
</style>
""", unsafe_allow_html=True)

# 로고와 제목
col1, col2 = st.columns([1, 5])
with col1:
    # 로고 표시
    logo_path = "pwc_logo.png"
    if os.path.exists(logo_path):
        st.image(logo_path, width=100)
    else:
        st.error("로고 파일을 찾을 수 없습니다. 프로젝트 루트에 'pwc_logo.png' 파일을 추가해주세요.")

with col2:
    st.markdown("<h1 class='main-title'>PwC 뉴스 분석기</h1>", unsafe_allow_html=True)
    st.markdown("회계법인 관점에서 중요한 뉴스를 자동으로 분석하는 AI 도구")

# 기본 선택 카테고리를 Anchor로 설정
COMPANIES = COMPANY_CATEGORIES["Anchor"]

# 사이드바 설정
st.sidebar.title("🔍 분석 설정")

# 0단계: 기본 설정
st.sidebar.markdown("### 📋 0단계: 기본 설정")

# 유효 언론사 설정
valid_press_dict = st.sidebar.text_area(
    "📰 유효 언론사 설정",
    value="""조선일보: ["조선일보", "chosun", "chosun.com"]
    중앙일보: ["중앙일보", "joongang", "joongang.co.kr", "joins.com"]
    동아일보: ["동아일보", "donga", "donga.com"]
    조선비즈: ["조선비즈", "chosunbiz", "biz.chosun.com"]
    매거진한경: ["매거진한경", "magazine.hankyung", "magazine.hankyung.com"]
    한국경제: ["한국경제", "한경", "hankyung", "hankyung.com", "한경닷컴"]
    매일경제: ["매일경제", "매경", "mk", "mk.co.kr"]
    연합뉴스: ["연합뉴스", "yna", "yna.co.kr"]
    파이낸셜뉴스: ["파이낸셜뉴스", "fnnews", "fnnews.com"]
    데일리팜: ["데일리팜", "dailypharm", "dailypharm.com"]
    IT조선: ["it조선", "it.chosun.com", "itchosun"]
    머니투데이: ["머니투데이", "mt", "mt.co.kr"]
    비즈니스포스트: ["비즈니스포스트", "businesspost", "businesspost.co.kr"]
    이데일리: ["이데일리", "edaily", "edaily.co.kr"]
    아시아경제: ["아시아경제", "asiae", "asiae.co.kr"]
    뉴스핌: ["뉴스핌", "newspim", "newspim.com"]
    뉴시스: ["뉴시스", "newsis", "newsis.com"]
    헤럴드경제: ["헤럴드경제", "herald", "heraldcorp", "heraldcorp.com"]""",
    help="분석에 포함할 신뢰할 수 있는 언론사와 그 별칭을 설정하세요. 형식: '언론사: [별칭1, 별칭2, ...]'",
    key="valid_press_dict"
)

# 추가 언론사 설정 (재평가 시에만 사용됨)
additional_press_dict = st.sidebar.text_area(
    "📰 추가 언론사 설정 (재평가 시에만 사용)",
    value="""철강금속신문: ["철강금속신문", "snmnews", "snmnews.com"]
    에너지신문: ["에너지신문", "energy-news", "energy-news.co.kr"]
    이코노믹데일리: ["이코노믹데일리", "economidaily", "economidaily.com"]""",
    help="기본 언론사에서 뉴스가 선택되지 않을 경우, 재평가 단계에서 추가로 고려할 언론사와 별칭을 설정하세요. 형식: '언론사: [별칭1, 별칭2, ...]'",
    key="additional_press_dict"
)

# 구분선 추가
st.sidebar.markdown("---")

# 날짜 필터 설정
st.sidebar.markdown("### 📅 날짜 필터")

# 현재 시간 가져오기
now = datetime.now()

# 기본 시작 날짜/시간 계산
default_start_date = now - timedelta(days=1)

# Set time to 8:00 AM for both start and end - 한국 시간 기준
start_datetime = datetime.combine(default_start_date.date(), 
                                    datetime.strptime("08:00", "%H:%M").time(), KST)
end_datetime = datetime.combine(now.date(), 
                                datetime.strptime("08:00", "%H:%M").time(), KST)

col1, col2 = st.sidebar.columns(2)
with col1:
    start_date = st.date_input(
        "시작 날짜",
        value=default_start_date.date(),
        help="이 날짜부터 뉴스를 검색합니다. 월요일인 경우 지난 금요일, 그 외에는 전일로 자동 설정됩니다."
    )
    start_time = st.time_input(
        "시작 시간",
        value=start_datetime.time(),
        help="시작 날짜의 구체적인 시간을 설정합니다. 기본값은 오전 8시입니다."
    )
with col2:
    end_date = st.date_input(
        "종료 날짜",
        value=now.date(),
        help="이 날짜까지의 뉴스를 검색합니다."
    )
    end_time = st.time_input(
        "종료 시간",
        value=end_datetime.time(),
        help="종료 날짜의 구체적인 시간을 설정합니다. 기본값은 오전 8시입니다."
    )

# 구분선 추가
st.sidebar.markdown("---")

# 1단계: 제외 판단 기준

# 기업 선택 섹션 제목
st.sidebar.markdown("### 🏢 분석할 기업 선택")

# 기업 카테고리 선택
selected_category = st.sidebar.radio(
    "기업 카테고리를 선택하세요",
    options=list(COMPANY_CATEGORIES.keys()),
    index=0,  # Anchor를 기본값으로 설정
    help="분석할 기업 카테고리를 선택하세요. Anchor(핵심), Growth(성장), Whitespace(신규) 중에서 선택할 수 있습니다."
)

# 선택된 카테고리에 따라 COMPANIES 업데이트
COMPANIES = COMPANY_CATEGORIES[selected_category]

# 새로운 기업 추가 섹션
new_company = st.sidebar.text_input(
    "새로운 기업 추가",
    value="",
    help="분석하고 싶은 기업명을 입력하고 Enter를 누르세요. (예: 네이버, 카카오, 현대중공업 등)"
)

# 새로운 기업 추가 로직 수정
if new_company and new_company not in COMPANIES:
    # 현재 선택된 카테고리에 기업 추가
    COMPANY_CATEGORIES[selected_category].append(new_company)
    # 세션 상태의 카테고리도 업데이트
    if 'company_categories' in st.session_state:
        st.session_state.company_categories[selected_category].append(new_company)
    # COMPANIES 리스트도 업데이트
    COMPANIES = COMPANY_CATEGORIES[selected_category]
    # 새 기업에 대한 기본 연관 키워드 설정 (기업명 자체만 포함)
    COMPANY_KEYWORD_MAP[new_company] = [new_company]
    # 세션 상태도 함께 업데이트
    if 'company_keyword_map' in st.session_state:
        st.session_state.company_keyword_map[new_company] = [new_company]

# 키워드 선택을 multiselect로 변경
selected_companies = st.sidebar.multiselect(
    "분석할 기업을 선택하세요 (최대 10개)",
    options=COMPANIES,
    default=COMPANIES[:10],  # 처음 10개 기업만 기본 선택으로 설정
    max_selections=10,
    help="분석하고자 하는 기업을 선택하세요. 한 번에 최대 10개까지 선택 가능합니다."
)

# 연관 키워드 관리 섹션
st.sidebar.markdown("### 🔍 연관 키워드 관리")
st.sidebar.markdown("각 기업의 연관 키워드를 확인하고 편집할 수 있습니다.")

# 세션 상태에 COMPANY_KEYWORD_MAP 및 COMPANY_CATEGORIES 저장 (초기화)
if 'company_keyword_map' not in st.session_state:
    st.session_state.company_keyword_map = COMPANY_KEYWORD_MAP.copy()
    
# 세션 상태에 회사 카테고리 저장 (초기화)
if 'company_categories' not in st.session_state:
    st.session_state.company_categories = COMPANY_CATEGORIES.copy()
else:
    # 세션에 저장된 카테고리 정보가 있으면 사용
    COMPANY_CATEGORIES = st.session_state.company_categories
    # 선택된 카테고리에 따라 COMPANIES 다시 업데이트
    COMPANIES = COMPANY_CATEGORIES[selected_category]

# 연관 키워드 UI 개선
if selected_companies:
    # 선택된 기업 중에서 관리할 기업 선택
    company_to_edit = st.sidebar.selectbox(
        "연관 키워드를 관리할 기업 선택",
        options=selected_companies,
        help="키워드를 확인하거나 추가할 기업을 선택하세요."
    )
    
    if company_to_edit:
        # 현재 연관 키워드 표시 (세션 상태에서 가져옴)
        current_keywords = st.session_state.company_keyword_map.get(company_to_edit, [company_to_edit])
        st.sidebar.markdown(f"**현재 '{company_to_edit}'의 연관 키워드:**")
        keyword_list = ", ".join(current_keywords)
        st.sidebar.code(keyword_list)
        
        # 연관 키워드 편집
        new_keywords = st.sidebar.text_area(
            "연관 키워드 편집",
            value=keyword_list,
            help="쉼표(,)로 구분하여 키워드를 추가/편집하세요.",
            key=f"edit_{company_to_edit}"  # 고유 키 추가
        )
        
        # 키워드 업데이트 함수
        def update_keywords():
            # 쉼표로 구분된 텍스트를 리스트로 변환
            updated_keywords = [kw.strip() for kw in new_keywords.split(",") if kw.strip()]
            
            # 업데이트
            if updated_keywords:
                st.session_state.company_keyword_map[company_to_edit] = updated_keywords
                st.sidebar.success(f"'{company_to_edit}'의 연관 키워드가 업데이트되었습니다!")
            else:
                # 비어있으면 기업명 자체만 포함
                st.session_state.company_keyword_map[company_to_edit] = [company_to_edit]
                st.sidebar.warning(f"연관 키워드가 비어있어 기업명만 포함됩니다.")
        
        # 변경 사항 적용 버튼
        if st.sidebar.button("연관 키워드 업데이트", key=f"update_{company_to_edit}", on_click=update_keywords):
            pass  # 실제 업데이트는 on_click에서 처리되므로 여기서는 아무것도 하지 않음

# 미리보기 버튼 - 모든 검색어 확인
with st.sidebar.expander("🔍 전체 검색 키워드 미리보기"):
    for i, company in enumerate(selected_companies, 1):
        # 세션 상태에서 키워드 가져오기
        company_keywords = st.session_state.company_keyword_map.get(company, [company])
        st.markdown(f"**{i}. {company}**")
        # 연관 키워드 표시
        for j, kw in enumerate(company_keywords, 1):
            st.write(f"  {j}) {kw}")

# 선택된 키워드들을 통합 (검색용)
keywords = []
for company in selected_companies:
    # 기업명 자체와 연관 키워드 모두 추가 (세션 상태에서 가져옴)
    company_keywords = st.session_state.company_keyword_map.get(company, [company])
    keywords.extend(company_keywords)

# 중복 제거
keywords = list(set(keywords))

# 구분선 추가
st.sidebar.markdown("---")

# 회사별 특화 기준 관리 섹션
st.sidebar.markdown("### 🎯 회사별 특화 기준 관리")
st.sidebar.markdown("각 기업의 AI 분석 특화 기준을 확인하고 편집할 수 있습니다.")

# 회사별 특화 기준 관리 UI
if selected_companies:
    # 선택된 기업 중에서 관리할 기업 선택
    company_to_manage = st.sidebar.selectbox(
        "특화 기준을 관리할 기업 선택",
        options=selected_companies,
        help="AI 분석 특화 기준을 확인하거나 편집할 기업을 선택하세요.",
        key="company_to_manage"
    )
    
    if company_to_manage:
        # 탭 형태로 1~3단계 기준을 구분
        criteria_tabs = st.sidebar.radio(
            f"'{company_to_manage}' 특화 기준 선택",
            ["1단계: 제외 기준", "2단계: 그룹핑 기준", "3단계: 선택 기준"],
            key=f"criteria_tabs_{company_to_manage}"
        )
        
        # 세션 상태에서 회사별 특화 기준 관리 (초기화)
        if 'company_additional_exclusion_criteria' not in st.session_state:
            st.session_state.company_additional_exclusion_criteria = COMPANY_ADDITIONAL_EXCLUSION_CRITERIA.copy()
        if 'company_additional_duplicate_handling' not in st.session_state:
            st.session_state.company_additional_duplicate_handling = COMPANY_ADDITIONAL_DUPLICATE_HANDLING.copy()
        if 'company_additional_selection_criteria' not in st.session_state:
            st.session_state.company_additional_selection_criteria = COMPANY_ADDITIONAL_SELECTION_CRITERIA.copy()
        
        if criteria_tabs == "1단계: 제외 기준":
            current_criteria = st.session_state.company_additional_exclusion_criteria.get(company_to_manage, "")
            st.sidebar.markdown(f"**현재 '{company_to_manage}'의 제외 특화 기준:**")
            if current_criteria.strip():
                st.sidebar.code(current_criteria, language="text")
            else:
                st.sidebar.info("설정된 특화 기준이 없습니다.")
            
            # 편집 영역
            new_exclusion_criteria = st.sidebar.text_area(
                "제외 특화 기준 편집",
                value=current_criteria,
                help="이 회사에만 적용될 추가 제외 기준을 입력하세요.",
                key=f"edit_exclusion_{company_to_manage}",
                height=150
            )
            
            # 업데이트 함수
            def update_exclusion_criteria():
                st.session_state.company_additional_exclusion_criteria[company_to_manage] = new_exclusion_criteria
                st.sidebar.success(f"'{company_to_manage}'의 제외 특화 기준이 업데이트되었습니다!")
            
            # 업데이트 버튼
            if st.sidebar.button("제외 기준 업데이트", key=f"update_exclusion_{company_to_manage}", on_click=update_exclusion_criteria):
                pass
                
        elif criteria_tabs == "2단계: 그룹핑 기준":
            current_criteria = st.session_state.company_additional_duplicate_handling.get(company_to_manage, "")
            st.sidebar.markdown(f"**현재 '{company_to_manage}'의 그룹핑 특화 기준:**")
            if current_criteria.strip():
                st.sidebar.code(current_criteria, language="text")
            else:
                st.sidebar.info("설정된 특화 기준이 없습니다.")
            
            # 편집 영역
            new_duplicate_criteria = st.sidebar.text_area(
                "그룹핑 특화 기준 편집",
                value=current_criteria,
                help="이 회사에만 적용될 추가 그룹핑 기준을 입력하세요.",
                key=f"edit_duplicate_{company_to_manage}",
                height=150
            )
            
            # 업데이트 함수
            def update_duplicate_criteria():
                st.session_state.company_additional_duplicate_handling[company_to_manage] = new_duplicate_criteria
                st.sidebar.success(f"'{company_to_manage}'의 그룹핑 특화 기준이 업데이트되었습니다!")
            
            # 업데이트 버튼
            if st.sidebar.button("그룹핑 기준 업데이트", key=f"update_duplicate_{company_to_manage}", on_click=update_duplicate_criteria):
                pass
                
        elif criteria_tabs == "3단계: 선택 기준":
            current_criteria = st.session_state.company_additional_selection_criteria.get(company_to_manage, "")
            st.sidebar.markdown(f"**현재 '{company_to_manage}'의 선택 특화 기준:**")
            if current_criteria.strip():
                st.sidebar.code(current_criteria, language="text")
            else:
                st.sidebar.info("설정된 특화 기준이 없습니다.")
            
            # 편집 영역
            new_selection_criteria = st.sidebar.text_area(
                "선택 특화 기준 편집",
                value=current_criteria,
                help="이 회사에만 적용될 추가 선택 기준을 입력하세요.",
                key=f"edit_selection_{company_to_manage}",
                height=150
            )
            
            # 업데이트 함수
            def update_selection_criteria():
                st.session_state.company_additional_selection_criteria[company_to_manage] = new_selection_criteria
                st.sidebar.success(f"'{company_to_manage}'의 선택 특화 기준이 업데이트되었습니다!")
            
            # 업데이트 버튼
            if st.sidebar.button("선택 기준 업데이트", key=f"update_selection_{company_to_manage}", on_click=update_selection_criteria):
                pass

# 미리보기 버튼 - 모든 회사별 특화 기준 확인
with st.sidebar.expander("🔍 전체 회사별 특화 기준 미리보기"):
    if selected_companies:
        # 세션 상태가 초기화되지 않은 경우를 위한 안전장치
        if 'company_additional_exclusion_criteria' not in st.session_state:
            st.session_state.company_additional_exclusion_criteria = COMPANY_ADDITIONAL_EXCLUSION_CRITERIA.copy()
        if 'company_additional_duplicate_handling' not in st.session_state:
            st.session_state.company_additional_duplicate_handling = COMPANY_ADDITIONAL_DUPLICATE_HANDLING.copy()
        if 'company_additional_selection_criteria' not in st.session_state:
            st.session_state.company_additional_selection_criteria = COMPANY_ADDITIONAL_SELECTION_CRITERIA.copy()
            
        for i, company in enumerate(selected_companies, 1):
            st.markdown(f"**{i}. {company}**")
            
            # 1단계 제외 기준 (세션 상태에서 가져오기)
            exclusion_criteria_text = st.session_state.company_additional_exclusion_criteria.get(company, "")
            if exclusion_criteria_text.strip():
                st.markdown("📝 **제외 특화 기준:**")
                st.text(exclusion_criteria_text[:100] + "..." if len(exclusion_criteria_text) > 100 else exclusion_criteria_text)
            
            # 2단계 그룹핑 기준 (세션 상태에서 가져오기)
            duplicate_criteria_text = st.session_state.company_additional_duplicate_handling.get(company, "")
            if duplicate_criteria_text.strip():
                st.markdown("🔄 **그룹핑 특화 기준:**")
                st.text(duplicate_criteria_text[:100] + "..." if len(duplicate_criteria_text) > 100 else duplicate_criteria_text)
            
            # 3단계 선택 기준 (세션 상태에서 가져오기)
            selection_criteria_text = st.session_state.company_additional_selection_criteria.get(company, "")
            if selection_criteria_text.strip():
                st.markdown("✅ **선택 특화 기준:**")
                st.text(selection_criteria_text[:100] + "..." if len(selection_criteria_text) > 100 else selection_criteria_text)
            
            if not (exclusion_criteria_text.strip() or duplicate_criteria_text.strip() or selection_criteria_text.strip()):
                st.info("설정된 특화 기준이 없습니다.")
            
            st.markdown("---")
    else:
        st.info("기업을 먼저 선택해주세요.")

# 구분선 추가
st.sidebar.markdown("---")

# GPT 모델 선택 섹션
st.sidebar.markdown("### 🤖 GPT 모델 선택")

selected_model = st.sidebar.selectbox(
    "분석에 사용할 GPT 모델을 선택하세요",
    options=list(GPT_MODELS.keys()),
    index=list(GPT_MODELS.keys()).index(DEFAULT_GPT_MODEL) if DEFAULT_GPT_MODEL in GPT_MODELS else 0,
    format_func=lambda x: f"{x} - {GPT_MODELS[x]}",
    help="각 모델의 특성:\n" + "\n".join([f"• {k}: {v}" for k, v in GPT_MODELS.items()])
)

# 모델 설명 표시
st.sidebar.markdown(f"""
<div style='background-color: #f0f2f6; padding: 10px; border-radius: 5px; margin-bottom: 20px;'>
    <strong>선택된 모델:</strong> {selected_model}<br>
    <strong>특징:</strong> {GPT_MODELS[selected_model]}
</div>
""", unsafe_allow_html=True)

# 구분선 추가
st.sidebar.markdown("---")

# 검색 결과 수 - 고정 값으로 설정
max_results = 100

# 시스템 프롬프트 설정
st.sidebar.markdown("### 🤖 시스템 프롬프트")

# 1단계: 제외 판단 시스템 프롬프트
system_prompt_1 = st.sidebar.text_area(
    "1단계: 제외 판단",
    value=SYSTEM_PROMPT_1,
    help="1단계 제외 판단에 사용되는 시스템 프롬프트를 설정하세요.",
    key="system_prompt_1",
    height=300
)

# 2단계: 그룹핑 시스템 프롬프트
system_prompt_2 = st.sidebar.text_area(
    "2단계: 그룹핑",
    value=SYSTEM_PROMPT_2,
    help="2단계 그룹핑에 사용되는 시스템 프롬프트를 설정하세요.",
    key="system_prompt_2",
    height=300
)

# 3단계: 중요도 평가 시스템 프롬프트
system_prompt_3 = st.sidebar.text_area(
    "3단계: 중요도 평가",
    value=SYSTEM_PROMPT_3,
    help="3단계 중요도 평가에 사용되는 시스템 프롬프트를 설정하세요.",
    key="system_prompt_3",
    height=300
)

st.sidebar.markdown("---")
st.sidebar.markdown("### 📋 1단계: 제외 판단 기준")

# 제외 기준 설정 - 기본 기준만 표시하고 사용자 수정 허용
exclusion_criteria = st.sidebar.text_area(
    "❌ 제외 기준",
    value=EXCLUSION_CRITERIA,
    help="분석에서 제외할 뉴스의 기준을 설정하세요. 실제 분석 시 각 회사별 특화 기준이 추가로 적용됩니다.",
    key="exclusion_criteria",
    height=300
)


# 구분선 추가
st.sidebar.markdown("---")

# 2단계: 그룹핑 기준
st.sidebar.markdown("### 📋 2단계: 그룹핑 기준")

# 중복 처리 기준 설정 - 기본 기준만 표시하고 사용자 수정 허용
duplicate_handling = st.sidebar.text_area(
    "🔄 중복 처리 기준",
    value=DUPLICATE_HANDLING,
    help="중복된 뉴스를 처리하는 기준을 설정하세요. 실제 분석 시 각 회사별 특화 기준이 추가로 적용됩니다.",
    key="duplicate_handling",
    height=300
)

# 구분선 추가
st.sidebar.markdown("---")

# 3단계: 선택 기준
st.sidebar.markdown("### 📋 3단계: 선택 기준")

# 선택 기준 설정 - 기본 기준만 표시하고 사용자 수정 허용
selection_criteria = st.sidebar.text_area(
    "✅ 선택 기준",
    value=SELECTION_CRITERIA,
    help="뉴스 선택에 적용할 주요 기준들을 나열하세요. 실제 분석 시 각 회사별 특화 기준이 추가로 적용됩니다.",
    key="selection_criteria",
    height=300
)

# 응답 형식 설정
response_format = st.sidebar.text_area(
    "📝 응답 형식",
    value="""선택된 뉴스 인덱스: [1, 3, 5]와 같은 형식으로 알려주세요.

각 선택된 뉴스에 대해:
제목: (뉴스 제목)
언론사: (언론사명)
발행일: (발행일자)
선정 사유: (구체적인 선정 이유)
분석 키워드: (해당 기업 그룹의 주요 계열사들)

[제외된 주요 뉴스]
제외된 중요 뉴스들에 대해:
인덱스: (뉴스 인덱스)
제목: (뉴스 제목)
제외 사유: (구체적인 제외 이유)""",
    help="분석 결과의 출력 형식을 설정하세요.",
    key="response_format",
    height=200
)

# 최종 프롬프트 생성
analysis_prompt = f"""
당신은 회계법인의 전문 애널리스트입니다. 아래 뉴스 목록을 분석하여 회계법인 관점에서 가장 중요한 뉴스를 선별하세요. 

[선택 기준]
{selection_criteria}

[제외 대상]
{exclusion_criteria}

[응답 요구사항]
1. 선택 기준에 부합하는 뉴스가 많다면 최대 3개까지 선택 가능합니다.
2. 선택 기준에 부합하는 뉴스가 없다면, 그 이유를 명확히 설명해주세요.

[응답 형식]
다음과 같은 JSON 형식으로 응답해주세요:

{{
    "selected_news": [
        {{
            "index": 1,
            "title": "뉴스 제목",
            "press": "언론사명",
            "date": "발행일자",
            "reason": "선정 사유",
            "keywords": ["키워드1", "키워드2"]
        }},
        ...
    ],
    "excluded_news": [
        {{
            "index": 2,
            "title": "뉴스 제목",
            "reason": "제외 사유"
        }},
        ...
    ]
}}

[유효 언론사]
{valid_press_dict}

[중복 처리 기준]
{duplicate_handling}
"""

# 메인 컨텐츠
if st.button("뉴스 분석 시작", type="primary"):
    # 이메일 미리보기를 위한 전체 내용 저장
    email_content = "[Client Intelligence]\n\n"
    
    # 모든 키워드 분석 결과를 저장할 딕셔너리
    all_results = {}
    
    for i, company in enumerate(selected_companies, 1):
        with st.spinner(f"'{company}' 관련 뉴스를 수집하고 분석 중입니다..."):
            # 해당 회사의 연관 키워드 확장 (세션 상태에서 가져옴)
            company_keywords = st.session_state.company_keyword_map.get(company, [company])
            
            # 연관 키워드 표시
            st.write(f"'{company}' 연관 키워드로 검색 중: {', '.join(company_keywords)}")
            
            # 사용자가 수정한 기준을 기본으로 하고, 해당 회사의 추가 특화 기준만 더함
            base_exclusion = exclusion_criteria
            base_duplicate = duplicate_handling
            base_selection = selection_criteria
            
            # 해당 회사의 추가 특화 기준만 가져오기 (세션 상태에서)
            # 세션 상태가 초기화되지 않은 경우를 위한 안전장치
            if 'company_additional_exclusion_criteria' not in st.session_state:
                st.session_state.company_additional_exclusion_criteria = COMPANY_ADDITIONAL_EXCLUSION_CRITERIA.copy()
            if 'company_additional_duplicate_handling' not in st.session_state:
                st.session_state.company_additional_duplicate_handling = COMPANY_ADDITIONAL_DUPLICATE_HANDLING.copy()
            if 'company_additional_selection_criteria' not in st.session_state:
                st.session_state.company_additional_selection_criteria = COMPANY_ADDITIONAL_SELECTION_CRITERIA.copy()
                
            company_additional_exclusion = st.session_state.company_additional_exclusion_criteria.get(company, "")
            company_additional_duplicate = st.session_state.company_additional_duplicate_handling.get(company, "")
            company_additional_selection = st.session_state.company_additional_selection_criteria.get(company, "")
            
            # 사용자 수정 기준 + 해당 회사 특화 기준 결합
            enhanced_exclusion_criteria = base_exclusion + company_additional_exclusion
            enhanced_duplicate_handling = base_duplicate + company_additional_duplicate  
            enhanced_selection_criteria = base_selection + company_additional_selection
            
            # initial_state 설정 부분 직전에 valid_press_dict를 딕셔너리로 변환하는 코드 추가
            # 텍스트 에어리어의 내용을 딕셔너리로 변환
            valid_press_config = {}
            try:
                # 문자열에서 딕셔너리 파싱
                lines = valid_press_dict.strip().split('\n')
                for line in lines:
                    line = line.strip()
                    if line and ': ' in line:
                        press_name, aliases_str = line.split(':', 1)
                        try:
                            # 문자열 형태의 리스트를 실제 리스트로 변환
                            aliases = eval(aliases_str.strip())
                            valid_press_config[press_name.strip()] = aliases
                            print(f"[DEBUG] Valid press 파싱 성공: {press_name.strip()} -> {aliases}")
                        except Exception as e:
                            print(f"[DEBUG] Valid press 파싱 실패: {line}, 오류: {str(e)}")
            except Exception as e:
                print(f"[DEBUG] Valid press 전체 파싱 실패: {str(e)}")
                # 오류 발생 시 빈 딕셔너리 사용
                valid_press_config = {}
            
            print(f"[DEBUG] 파싱된 valid_press_dict: {valid_press_config}")
            
            # 추가 언론사도 파싱
            additional_press_config = {}
            try:
                # 문자열에서 딕셔너리 파싱
                lines = additional_press_dict.strip().split('\n')
                for line in lines:
                    line = line.strip()
                    if line and ': ' in line:
                        press_name, aliases_str = line.split(':', 1)
                        try:
                            # 문자열 형태의 리스트를 실제 리스트로 변환
                            aliases = eval(aliases_str.strip())
                            additional_press_config[press_name.strip()] = aliases
                            print(f"[DEBUG] Additional press 파싱 성공: {press_name.strip()} -> {aliases}")
                        except Exception as e:
                            print(f"[DEBUG] Additional press 파싱 실패: {line}, 오류: {str(e)}")
            except Exception as e:
                print(f"[DEBUG] Additional press 전체 파싱 실패: {str(e)}")
                # 오류 발생 시 빈 딕셔너리 사용
                additional_press_config = {}
            
            print(f"[DEBUG] 파싱된 additional_press_dict: {additional_press_config}")
            
            # 각 키워드별 상태 초기화
            initial_state = {
                "news_data": [], 
                "filtered_news": [], 
                "analysis": "", 
                "keyword": company_keywords,  # 회사별 확장 키워드 리스트 전달
                "model": selected_model,
                "excluded_news": [],
                "borderline_news": [],
                "retained_news": [],
                "grouped_news": [],
                "final_selection": [],
                # 회사별 enhanced 기준들 적용
                "exclusion_criteria": enhanced_exclusion_criteria,
                "duplicate_handling": enhanced_duplicate_handling,
                "selection_criteria": enhanced_selection_criteria,
                "system_prompt_1": system_prompt_1,
                "user_prompt_1": "",
                "llm_response_1": "",
                "system_prompt_2": system_prompt_2,
                "user_prompt_2": "",
                "llm_response_2": "",
                "system_prompt_3": system_prompt_3,
                "user_prompt_3": "",
                "llm_response_3": "",
                "not_selected_news": [],
                "original_news_data": [],
                # 언론사 설정 추가 (파싱된 딕셔너리 사용)
                "valid_press_dict": valid_press_config,
                # 추가 언론사 설정 추가
                "additional_press_dict": additional_press_config,
                # 날짜 필터 정보 추가
                "start_datetime": datetime.combine(start_date, start_time, KST),
                "end_datetime": datetime.combine(end_date, end_time, KST)
                #"start_datetime": start_datetime,
                #"end_datetime": end_datetime
            }
            
            
            print(f"[DEBUG] start_datetime: {datetime.combine(start_date, start_time)}")
            print(f"[DEBUG] end_datetime: {datetime.combine(end_date, end_time)}")
            
            # 1단계: 뉴스 수집
            st.write("1단계: 뉴스 수집 중...")
            state_after_collection = collect_news(initial_state)
            
            # 2단계: 유효 언론사 필터링
            st.write("2단계: 유효 언론사 필터링 중...")
            state_after_press_filter = filter_valid_press(state_after_collection)
            
            # 3단계: 제외 판단
            st.write("3단계: 제외 판단 중...")
            state_after_exclusion = filter_excluded_news(state_after_press_filter)
            
            # 4단계: 그룹핑
            st.write("4단계: 그룹핑 중...")
            state_after_grouping = group_and_select_news(state_after_exclusion)
            
            # 5단계: 중요도 평가
            st.write("5단계: 중요도 평가 중...")
            final_state = evaluate_importance(state_after_grouping)

            # 6단계: 0개 선택 시 재평가 (개선된 코드)
            if len(final_state["final_selection"]) == 0:
                st.write("6단계: 선택된 뉴스가 없어 재평가를 시작합니다...")
                
                # 추가 언론사 설정 불러오기 (이미 파싱된 딕셔너리 사용)
                additional_press = additional_press_config
                
                # 기존 유효 언론사에 추가 언론사 병합 (딕셔너리 병합)
                expanded_valid_press_dict = {**valid_press_config, **additional_press}
                
                # 추가 언론사로 필터링한 뉴스 저장 (기존 뉴스와 구분)
                additional_valid_news = []
                
                # 확장된 언론사 목록으로 원본 뉴스 재필터링
                try:
                    # 현재 필터링된 유효 언론사 뉴스 수집
                    current_news_data = final_state.get("news_data", [])
                    
                    # 원본 뉴스 데이터 가져오기
                    original_news_data = final_state.get("original_news_data", [])
                    
                    if expanded_valid_press_dict:
                        # 확장된 언론사 목록으로 원본 뉴스 재필터링
                        for news in original_news_data:
                            # 이미 필터링된 뉴스는 제외
                            if any(existing_news.get('url') == news.get('url') for existing_news in current_news_data):
                                continue
                                
                            press = news.get("press", "").lower()
                            url = news.get("url", "").lower()
                            
                            # 추가된 언론사 기준으로만 필터링
                            is_valid = False
                            for main_press, aliases in expanded_valid_press_dict.items():
                                domain = urlparse(url).netloc.lower()
                                # 더 유연한 매칭 적용
                                if any(alias.lower() in press or press in alias.lower() for alias in aliases) or \
                                   any(alias.lower() in domain or domain in alias.lower() for alias in aliases):
                                    is_valid = True
                                    break
                            
                            if is_valid:
                                # 새 언론사 필터링된 뉴스임을 표시
                                additional_valid_news.append(news)
                    
                    # 추가 유효 뉴스가 있으면 기존 news_data에 추가
                    if additional_valid_news:
                        st.success(f"추가 언론사 기준으로 {len(additional_valid_news)}개의 뉴스가 추가로 필터링되었습니다.")
                        
                        # 기존 뉴스 데이터와 병합
                        combined_news = current_news_data + additional_valid_news
                        reevaluation_state = final_state.copy()
                        reevaluation_state["news_data"] = combined_news
                        
                        # 추가된 뉴스들에 대한 제외/유지 판단 재실행
                        reevaluation_state = filter_excluded_news(reevaluation_state)
                        
                        # 그룹핑 재실행
                        reevaluation_state = group_and_select_news(reevaluation_state)
                    else:
                        # 추가 뉴스가 없으면 원래 상태 복사
                        reevaluation_state = final_state.copy()
                        combined_news = current_news_data
                except Exception as e:
                    st.warning(f"추가 언론사 필터링 중 오류 발생: {str(e)}")
                    reevaluation_state = final_state.copy()
                    combined_news = final_state.get("news_data", [])
                
                # 확장된 유효 언론사 목록 문자열로 변환 (프롬프트용)
                expanded_valid_press_str = "유효 언론사 목록:\n"
                for press, aliases in expanded_valid_press_dict.items():
                    expanded_valid_press_str += f"  * {press}: {aliases}\n"
                
                # 재평가 시스템 프롬프트 개선 - 모든 뉴스 데이터 포함
                reevaluation_system_prompt = f"""
                당신은 회계법인의 뉴스 분석 전문가입니다. 현재 선정된 뉴스가 없어 재평가가 필요합니다.
                아래 4가지 방향으로 뉴스를 재검토하세요:

                1. 언론사 필터링 기준 완화:
                - 기존 유효 언론사 목록 외에도 다음 언론사의 기사를 포함하여 평가합니다:
                  * 철강금속신문: 산업 전문지로 금속/철강 업계 소식에 특화됨
                  * 에너지신문: 에너지 산업 전문 매체로 관련 기업 분석에 유용함
                  * 이코노믹데일리: 경제 전문지로 추가적인 시각 제공

                2. 제외 조건 재평가:
                - 제외 기준을 유연하게 적용하여, 회계법인의 관점에서 재무적 관점으로 해석 가능한 기사들을 보류로 분류
                - 특히 기업의 재정 혹은 전략적 변동과 연관된 기사를 보류로 전환

                3. 중복 제거 재평가:
                - 중복 기사 중에서도 언론사의 신뢰도나 기사 내용을 추가로 고려하여 가능한 경우 추가적으로 선택
                - 재무적/전략적 관점에서 추가 정보를 제공하는 기사 우선 선택

                4. 중요도 재평가:
                - 선택 기준을 일부 충족하지 않는 기사일지라도 기업명과 관련된 재정적 또는 전략적 변동에 대해서는 중요도를 '중'으로 평가
                - 필요하다면 중요도 '하'도 고려하여 최소 2개의 기사를 선정

                [확장된 유효 언론사 목록]
                {expanded_valid_press_str}

                [기존 제외 기준]
                {enhanced_exclusion_criteria}

                [기존 중복 처리 기준]
                {enhanced_duplicate_handling}

                [기존 선택 기준]
                {enhanced_selection_criteria}

                [전체 뉴스 목록]
                """
                
                # 모든 뉴스 데이터를 하나의 리스트로 통합 (JSON 형식으로)
                all_news_json = []
                for i, news in enumerate(combined_news):
                    all_news_json.append({
                        "index": i+1,
                        "title": news.get('content', '제목 없음'),
                        "url": news.get('url', ''),
                        "date": news.get('date', ''),
                        "press": news.get('press', '')
                    })
                
                # 프롬프트에 통합된 뉴스 목록 추가
                reevaluation_system_prompt += str(all_news_json)
                
                reevaluation_system_prompt += """
                
                [분류된 뉴스 목록]
                - 제외된 뉴스: {[f"제목: {news['title']}, 인덱스: {news['index']}, 사유: {news.get('reason', '')}" for news in reevaluation_state["excluded_news"]]}
                - 보류 뉴스: {[f"제목: {news['title']}, 인덱스: {news['index']}, 사유: {news.get('reason', '')}" for news in reevaluation_state["borderline_news"]]}
                - 유지 뉴스: {[f"제목: {news['title']}, 인덱스: {news['index']}, 사유: {news.get('reason', '')}" for news in reevaluation_state["retained_news"]]}

                ⚠️ 매우 중요한 지시사항 ⚠️
                1. 반드시 최소 2개 이상의 기사를 선정해야 합니다.
                2. 언론사와 기사 내용을 고려하여 선정 기준을 대폭 완화하세요.
                3. 원래 '제외'로 분류했던 기사 중에서도 회계법인 관점에서 조금이라도 가치가 있는 내용이 있다면 재검토하세요.
                4. 어떤 경우에도 2개 미만의 기사를 선정하지 마세요. 이는 절대적인 요구사항입니다.
                5. 모든 기사가 부적합하다고 판단되더라도 그 중에서 가장 나은 2개는 선정해야 합니다.
                6. 추가 언론사 목록의 기사들도 동등하게 고려하세요.

                다음 JSON 형식으로 응답해주세요:
                {
                    "reevaluated_news": [
                        {
                            "index": 1,
                            "title": "뉴스 제목",
                            "press": "언론사명",
                            "date": "발행일자",
                            "reason": "선정 사유",
                            "keywords": ["키워드1", "키워드2"],
                            "affiliates": ["계열사1", "계열사2"],
                            "importance": "중요도(상/중/하)"
                        }
                    ]
                }
                """
                
                # 재평가 시스템 프롬프트로 업데이트
                reevaluation_state["system_prompt_3"] = reevaluation_system_prompt
                
                # 재평가 실행 (evaluate_importance 함수 재사용)
                st.write("- 제외/중복/중요도 통합 재평가 중...")
                reevaluation_result = evaluate_importance(reevaluation_state)
                
                # 재평가 결과가 있으면 최종 상태 업데이트
                if "final_selection" in reevaluation_result and reevaluation_result["final_selection"]:
                    final_state["final_selection"] = reevaluation_result["final_selection"]
                    # 재평가 결과임을 표시하기 위한 필드 추가
                    final_state["is_reevaluated"] = True
                    st.success(f"재평가 후 {len(final_state['final_selection'])}개의 뉴스가 선택되었습니다.")
                else:
                    # 그래도 없으면 오류 메시지만 표시
                    st.error("재평가 후에도 선정할 수 있는 뉴스가 없습니다.")

            # 키워드별 분석 결과 저장
            all_results[company] = final_state["final_selection"]
            
            # 키워드 구분선 추가
            st.markdown("---")
            
            # 키워드별 섹션 구분
            st.markdown(f"## 📊 {company} 분석 결과")
            
            # 전체 뉴스 표시 (필터링 전)
            with st.expander(f"📰 '{company}' 관련 전체 뉴스 (필터링 전)"):
                for i, news in enumerate(final_state.get("original_news_data", []), 1):
                    date_str = news.get('date', '날짜 정보 없음')
                    url = news.get('url', 'URL 정보 없음')
                    press = news.get('press', '알 수 없음')
                    st.markdown(f"""
                    <div class="news-card">
                        <div class="news-title">{i}. {news['content']}</div>
                        <div class="news-meta">📰 {press}</div>
                        <div class="news-date">📅 {date_str}</div>
                        <div class="news-url">🔗 <a href="{url}" target="_blank">{url}</a></div>
                    </div>
                    """, unsafe_allow_html=True)
            
            # 유효 언론사 필터링된 뉴스 표시
            with st.expander(f"📰 '{company}' 관련 유효 언론사 뉴스"):
                for i, news in enumerate(final_state["news_data"]):
                    date_str = news.get('date', '날짜 정보 없음')
                    url = news.get('url', 'URL 정보 없음')
                    press = news.get('press', '알 수 없음')
                    st.markdown(f"""
                    <div class="news-card">
                        <div class="news-title">{i+1}. {news['content']}</div>
                        <div class="news-meta">📰 {press}</div>
                        <div class="news-date">📅 {date_str}</div>
                        <div class="news-url">🔗 <a href="{url}" target="_blank">{url}</a></div>
                    </div>
                    """, unsafe_allow_html=True)
            
            # 2단계: 유효 언론사 필터링 결과 표시
            st.markdown("<div class='subtitle'>🔍 2단계: 유효 언론사 필터링 결과</div>", unsafe_allow_html=True)
            st.markdown(f"유효 언론사 뉴스: {len(final_state['news_data'])}개")
            
            # 3단계: 제외/보류/유지 뉴스 표시
            st.markdown("<div class='subtitle'>🔍 3단계: 뉴스 분류 결과</div>", unsafe_allow_html=True)
            
            # 제외된 뉴스
            with st.expander("❌ 제외된 뉴스"):
                for news in final_state["excluded_news"]:
                    st.markdown(f"<div class='excluded-news'>[{news['index']}] {news['title']}<br/>└ {news['reason']}</div>", unsafe_allow_html=True)
            
            # 보류 뉴스
            with st.expander("⚠️ 보류 뉴스"):
                for news in final_state["borderline_news"]:
                    st.markdown(f"<div class='excluded-news'>[{news['index']}] {news['title']}<br/>└ {news['reason']}</div>", unsafe_allow_html=True)
            
            # 유지 뉴스
            with st.expander("✅ 유지 뉴스"):
                for news in final_state["retained_news"]:
                    st.markdown(f"<div class='excluded-news'>[{news['index']}] {news['title']}<br/>└ {news['reason']}</div>", unsafe_allow_html=True)
            
            # 4단계: 그룹핑 결과 표시
            st.markdown("<div class='subtitle'>🔍 4단계: 뉴스 그룹핑 결과</div>", unsafe_allow_html=True)
            
            with st.expander("📋 그룹핑 결과 보기"):
                for group in final_state["grouped_news"]:
                    st.markdown(f"""
                    <div class="analysis-section">
                        <h4>그룹 {group['indices']}</h4>
                        <p>선택된 기사: {group['selected_index']}</p>
                        <p>선정 이유: {group['reason']}</p>
                    </div>
                    """, unsafe_allow_html=True)
            
            # 5단계: 최종 선택 결과 표시
            st.markdown("<div class='subtitle'>🔍 5단계: 최종 선택 결과</div>", unsafe_allow_html=True)
            
            # 재평가 여부 확인 (is_reevaluated 필드 있으면 재평가된 것)
            was_reevaluated = final_state.get("is_reevaluated", False)
            
            # 재평가 여부에 따라 메시지와 스타일 변경
            if was_reevaluated:
                # 재평가가 수행된 경우 6단계 표시
                st.warning("5단계에서 선정된 뉴스가 없어 6단계 재평가를 진행했습니다.")
                st.markdown("<div class='subtitle'>🔍 6단계: 재평가 결과</div>", unsafe_allow_html=True)
                st.markdown("### 📰 재평가 후 선정된 뉴스")
                # 재평가 스타일 적용
                news_style = "border-left: 4px solid #FFA500; background-color: #FFF8DC;"
                reason_prefix = "<span style=\"color: #FFA500; font-weight: bold;\">재평가 후</span> 선별 이유: "
            else:
                # 정상적으로 5단계에서 선정된 경우
                st.markdown("### 📰 최종 선정된 뉴스")  
                # 일반 스타일 적용
                news_style = ""
                reason_prefix = "선별 이유: "
            
            # 최종 선정된 뉴스 표시
            for news in final_state["final_selection"]:
                # 날짜 형식 변환
                
                date_str = format_date(news.get('date', ''))
                
                try:
                    # YYYY-MM-DD 형식으로 가정
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%m/%d')
                except Exception as e:
                    try:
                        # GMT 형식 시도
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%m/%d')
                    except Exception as e:
                        formatted_date = date_str if date_str else '날짜 정보 없음'

                url = news.get('url', 'URL 정보 없음')
                press = news.get('press', '언론사 정보 없음')
                
                # 뉴스 정보 표시
                st.markdown(f"""
                    <div class="selected-news" style="{news_style}">
                        <div class="news-title-large">{news['title']} ({formatted_date})</div>
                        <div class="news-url">🔗 <a href="{url}" target="_blank">{url}</a></div>
                        <div class="selection-reason">
                            • {reason_prefix}{news['reason']}
                        </div>
                        <div class="news-summary">
                            • 키워드: {', '.join(news['keywords'])} | 관련 계열사: {', '.join(news['affiliates'])} | 언론사: {press}
                        </div>
                    </div>
                """, unsafe_allow_html=True)
                
                # 구분선 추가
                st.markdown("---")
            
            # 선정되지 않은 뉴스 표시
            if final_state.get("not_selected_news"):
                with st.expander("❌ 선정되지 않은 뉴스"):
                    for news in final_state["not_selected_news"]:
                        st.markdown(f"""
                        <div class="not-selected-news">
                            <div class="news-title">{news['index']}. {news['title']}</div>
                            <div class="importance-low">💡 중요도: {news['importance']}</div>
                            <div class="not-selected-reason">❌ 미선정 사유: {news['reason']}</div>
                        </div>
                        """, unsafe_allow_html=True)
            
            # 디버그 정보
            with st.expander("디버그 정보"):
                st.markdown("### 1단계: 제외 판단")
                st.markdown("#### 시스템 프롬프트")
                st.text(final_state.get("system_prompt_1", "없음"))
                st.markdown("#### 사용자 프롬프트")
                st.text(final_state.get("user_prompt_1", "없음"))
                st.markdown("#### LLM 응답")
                st.text(final_state.get("llm_response_1", "없음"))
                
                st.markdown("### 2단계: 그룹핑")
                st.markdown("#### 시스템 프롬프트")
                st.text(final_state.get("system_prompt_2", "없음"))
                st.markdown("#### 사용자 프롬프트")
                st.text(final_state.get("user_prompt_2", "없음"))
                st.markdown("#### LLM 응답")
                st.text(final_state.get("llm_response_2", "없음"))
                
                st.markdown("### 3단계: 중요도 평가")
                st.markdown("#### 시스템 프롬프트")
                st.text(final_state.get("system_prompt_3", "없음"))
                st.markdown("#### 사용자 프롬프트")
                st.text(final_state.get("user_prompt_3", "없음"))
                st.markdown("#### LLM 응답")
                st.text(final_state.get("llm_response_3", "없음"))
                
                # 6단계: 재평가 정보 추가
                if final_state.get("is_reevaluated", False):
                    st.markdown("### 4단계: 재평가")
                    st.markdown("#### 시스템 프롬프트")
                    # 실제 사용된 재평가 시스템 프롬프트 표시
                    st.text(reevaluation_state.get("system_prompt_3", "없음") if 'reevaluation_state' in locals() else "재평가 프롬프트 정보 없음")
                    st.markdown("#### 사용자 프롬프트")
                    st.text(reevaluation_state.get("user_prompt_3", "없음") if 'reevaluation_state' in locals() else "재평가 사용자 프롬프트 정보 없음")
                    st.markdown("#### LLM 응답")
                    st.text(reevaluation_state.get("llm_response_3", "없음") if 'reevaluation_state' in locals() else "재평가 LLM 응답 정보 없음")
            
            # 이메일 내용 추가
            email_content += f"{i}. {company}\n"
            for news in final_state["final_selection"]:
                # 날짜 형식 변환
                date_str = news.get('date', '')
                try:
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%m/%d')
                except Exception as e:
                    try:
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%m/%d')
                    except Exception as e:
                        formatted_date = date_str if date_str else '날짜 정보 없음'
                
                url = news.get('url', '')
                email_content += f"  - {news['title']} ({formatted_date}) {url}\n"
            email_content += "\n"
            
            # 키워드 구분선 추가
            st.markdown("---")

    # 모든 키워드 분석이 끝난 후 이메일 미리보기 섹션 추가
    st.markdown("<div class='subtitle'>📧 이메일 미리보기</div>", unsafe_allow_html=True)
    
    # HTML 버전 생성
    html_email_content = "<div style='font-family: Arial, sans-serif; max-width: 800px; font-size: 14px; line-height: 1.5;'>"
    
    html_email_content += "<div style='margin-top: 20px; font-size: 14px;'>안녕하세요, 좋은 아침입니다!<br>오늘의 Client Intelligence 전달 드립니다.<br><br></div>"
    plain_email_content = "\n안녕하세요, 좋은 아침입니다!\n오늘의 Client Intelligence 전달 드립니다."
    
    html_email_content += "<div style='font-size: 14px; font-weight: bold; margin-bottom: 15px; border-bottom: 1px solid #000;'>[Client Intelligence]</div>"
    
    # 일반 텍스트 버전 생성 (복사용)
    plain_email_content += "[Client Intelligence]\n\n"
    
    def clean_title(title):
        """Clean title by removing the press name pattern at the end"""
        # Remove the press pattern (e.g., '제목 - 조선일보', '제목-조선일보', '제목 - Chosun Biz')
        title = re.sub(r"\s*-\s*[가-힣A-Za-z0-9\s]+$", "", title).strip()
        return title

    for i, company in enumerate(selected_companies, 1):
        # HTML 버전에서 키워드를 파란색으로 표시
        html_email_content += f"<div style='font-size: 14px; font-weight: bold; margin-top: 15px; margin-bottom: 10px; color: #0000FF;'>{i}. {company}</div>"
        html_email_content += "<ul style='list-style-type: none; padding-left: 20px; margin: 0;'>"
        
        # 텍스트 버전에서도 키워드 구분을 위해 줄바꿈 추가
        plain_email_content += f"{i}. {company}\n"
        
        # 해당 키워드의 뉴스 가져오기
        news_list = all_results.get(company, [])
        
        if not news_list:
            # 최종 선정 뉴스가 0건인 경우 안내 문구 추가
            html_email_content += "<li style='margin-bottom: 8px; font-size: 14px; color: #888;'>AI 분석결과 금일자로 회계법인 관점에서 특별히 주목할 만한 기사가 없습니다.</li>"
            plain_email_content += "  - AI 분석결과 금일자로 회계법인 관점에서 특별히 주목할 만한 기사가 없습니다.\n"
        else:
            for news in news_list:
                # 날짜 형식 변환
                date_str = news.get('date', '')
                try:
                    date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                    formatted_date = date_obj.strftime('%m/%d')
                except Exception as e:
                    try:
                        date_obj = datetime.strptime(date_str, '%a, %d %b %Y %H:%M:%S %Z')
                        formatted_date = date_obj.strftime('%m/%d')
                    except Exception as e:
                        formatted_date = date_str if date_str else '날짜 정보 없음'
                
                url = news.get('url', '')
                title = news.get('title', '')
                # 이메일 미리보기에서는 언론사 패턴 제거
                title = clean_title(title)
                # HTML 버전 - 링크를 [파일 링크]로 표시하고 글자 크기 통일, 본문 bold 처리
                html_email_content += f"<li style='margin-bottom: 8px; font-size: 14px;'><span style='font-weight: bold;'>- {title} ({formatted_date})</span> <a href='{url}' style='color: #1a0dab; text-decoration: none;'>[기사 링크]</a></li>"
                
                # 텍스트 버전 - 링크를 [파일 링크]로 표시하고 실제 URL은 그 다음 줄에
                plain_email_content += f"  - {title} ({formatted_date}) [기사 링크]\n    {url}\n"
        
        html_email_content += "</ul>"
        plain_email_content += "\n"
    
    # 서명 추가
    html_email_content += "<div style='margin-top: 20px; font-size: 14px;'><br>감사합니다.<br>Client & Market 드림</div>"
    plain_email_content += "\n감사합니다.\nClient & Market 드림"
    
    html_email_content += "</div>"
    
    # 이메일 미리보기 표시
    st.markdown(f"<div class='email-preview'>{html_email_content}</div>", unsafe_allow_html=True)



else:
    # 초기 화면 설명 (주석 처리됨)
    """
    ### 👋 PwC 뉴스 분석기에 오신 것을 환영합니다!
    
    이 도구는 입력한 키워드에 대한 최신 뉴스를 자동으로 수집하고, 회계법인 관점에서 중요한 뉴스를 선별하여 분석해드립니다.
    
    #### 주요 기능:
    1. 최신 뉴스 자동 수집 (기본 100개)
    2. 신뢰할 수 있는 언론사 필터링
    3. 6단계 AI 기반 뉴스 분석 프로세스:
       - 1단계: 뉴스 수집 - 키워드 기반으로 최신 뉴스 데이터 수집
       - 2단계: 유효 언론사 필터링 - 신뢰할 수 있는 언론사 선별
       - 3단계: 제외/보류/유지 판단 - 회계법인 관점에서의 중요도 1차 분류
       - 4단계: 유사 뉴스 그룹핑 - 중복 기사 제거 및 대표 기사 선정
       - 5단계: 중요도 평가 및 최종 선정 - 회계법인 관점의 중요도 평가
       - 6단계: 필요시 재평가 - 선정된 뉴스가 없을 경우 AI가 기준을 완화하여 재평가
    4. 선별된 뉴스에 대한 상세 정보 제공
       - 제목 및 날짜
       - 원문 링크
       - 선별 이유
       - 키워드, 관련 계열사, 언론사 정보
    5. 분석 결과 이메일 형식 미리보기
    
    #### 사용 방법:
    1. 사이드바에서 분석할 기업을 선택하세요 (최대 10개)
       - 기본 제공 기업 목록에서 선택
       - 새로운 기업 직접 추가 가능
    2. GPT 모델을 선택하세요
       - gpt-4o: 빠르고 실시간 (기본값)
    3. 날짜 필터를 설정하세요
       - 기본값: 어제 또는 지난 금요일(월요일인 경우)부터 오늘까지
    4. "뉴스 분석 시작" 버튼을 클릭하세요
    
    #### 분석 결과 확인:
    - 각 키워드별 최종 선정된 중요 뉴스
    - 선정 과정의 중간 결과(제외/보류/유지, 그룹핑 등)
    - 선정된 모든 뉴스의 요약 이메일 미리보기
    - 디버그 정보 (시스템 프롬프트, AI 응답 등)
    
    """

# 푸터
st.markdown("---")
st.markdown("© 2024 PwC 뉴스 분석기 | 회계법인 관점의 뉴스 분석 도구")
